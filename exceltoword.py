import tempfile
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import xlwings as xw
from docx import Document
from docx.shared import Cm, RGBColor
from docx.enum.text import WD_BREAK
import re
import time
from PIL import Image
import os
import sys ; sys.setrecursionlimit(sys.getrecursionlimit() * 5)

# Find poppler-mappen, uanset om scriptet kører som .py eller .exe
if getattr(sys, 'frozen', False):
    script_dir = sys._MEIPASS  # PyInstaller bruger _MEIPASS som midlertidig mappe
else:
    script_dir = os.path.dirname(os.path.abspath(sys.argv[0])) if "__file__" in globals() else os.getcwd()

poppler_path = os.path.join(script_dir, "poppler", "Library", "bin")

# Brug Poppler-stien i pdf2image
from pdf2image import convert_from_path


def select_excel_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx;*.xlsm")])
    if file_path:
        excel_entry.delete(0, tk.END)
        excel_entry.insert(0, file_path)
        load_sheets(file_path)


def load_sheets(file_path):
    try:
        with xw.App(visible=False) as app:
            wb = app.books.open(file_path)
            sheet_names = [sheet.name for sheet in wb.sheets if sheet.visible]

        # Opdater GUI med de synlige ark
        for widget in sheets_frame.winfo_children():
            widget.destroy()

        # Opret Canvas til scrollbar
        canvas = tk.Canvas(sheets_frame)
        scrollbar = tk.Scrollbar(sheets_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Tilføj musescroll-funktion
        def on_mouse_wheel(event):
            canvas.yview_scroll(-1 * (event.delta // 120), "units")

        canvas.bind_all("<MouseWheel>", on_mouse_wheel)

        # Placer canvas og scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Tilføj checkbokse til arkene
        global sheet_checkbuttons
        sheet_checkbuttons = []
        for sheet in sheet_names:
            var = tk.BooleanVar()
            chk = tk.Checkbutton(scrollable_frame, text=sheet, variable=var)
            chk.var = var
            chk.pack(anchor="w")
            sheet_checkbuttons.append(chk)

    except Exception as e:
        messagebox.showerror("Fejl", f"Kunne ikke indlæse ark: {e}")
        return


def select_word_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    if file_path:
        word_entry.delete(0, tk.END)
        word_entry.insert(0, file_path)


def select_image_output_folder():
    folder_path = filedialog.askdirectory()
    if folder_path:
        output_entry.delete(0, tk.END)
        output_entry.insert(0, folder_path)


def generate_unique_filename(base_path):
    counter = 1
    file_path = f"{base_path}_opdateret.docx"
    while os.path.exists(file_path):
        file_path = f"{base_path}_opdateret_{counter}.docx"
        counter += 1
    return file_path


def crop_image(image_path):
    try:
        img = Image.open(image_path)
        dpi = 300  # Samme dpi som ved konvertering

        # Læs værdier fra GUI'en
        top_crop = int(float(top_crop_entry.get()) / 25.4 * dpi)
        bottom_crop = int(float(bottom_crop_entry.get()) / 25.4 * dpi)
        left_crop = int(float(left_crop_entry.get()) / 25.4 * dpi)
        right_crop = int(float(right_crop_entry.get()) / 25.4 * dpi)

        width, height = img.size
        cropped_img = img.crop((left_crop, top_crop, width - right_crop, height - bottom_crop))
        cropped_img.save(image_path)
    except Exception as e:
        messagebox.showerror("Fejl i beskæring", f"Beskæring fejlede for {os.path.basename(image_path)}:\n{e}")


def save_excel_sheets_as_png():
    excel_file = excel_entry.get()
    output_folder = output_entry.get()
    selected_sheets = [chk.cget("text") for chk in sheet_checkbuttons if chk.var.get()]

    if not excel_file or not selected_sheets or not output_folder:
        messagebox.showwarning("Fejl", "Vælg Excel-fil, ark og billedoutput-mappe!")
        return

    try:
        with xw.App(visible=False) as app:
            wb = app.books.open(os.path.abspath(excel_file))
            total_sheets = len(selected_sheets)
            start_time = time.time()

            for index, sheet_name in enumerate(selected_sheets, start=1):
                sheet = wb.sheets[sheet_name]
                temp_pdf = os.path.join(tempfile.gettempdir(), f"{sheet_name}.pdf")

                try:
                    sheet.to_pdf(temp_pdf)
                    images = convert_from_path(temp_pdf, dpi=300, poppler_path=poppler_path)

                    for i, img in enumerate(images, start=1):
                        img_path = os.path.join(output_folder, f"{sheet_name}_{i}.png")
                        img.save(img_path, "PNG")
                        crop_image(img_path)  # Beskær billedet
                except Exception as e:
                    messagebox.showerror("Fejl", f"Kunne ikke eksportere {sheet_name}: {e}")

                elapsed = time.time() - start_time
                remaining = (elapsed / index) * (total_sheets - index)
                progress_var.set(int((index / total_sheets) * 100))
                progress_label.config(text=f"Forventet tid tilbage: {int(remaining)} sekunder")
                root.update_idletasks()

            messagebox.showinfo("Succes", "Billeder gemt!")
    except Exception as e:
        messagebox.showerror("Fejl", f"Kunne ikke åbne Excel: {e}")


def insert_images_into_word():
    word_file = word_entry.get()
    output_folder = output_entry.get()

    if not word_file or not output_folder:
        messagebox.showwarning("Fejl", "Vælg Word-fil og billedoutput-mappe!")
        return

    base_name, _ = os.path.splitext(word_file)
    output_file = generate_unique_filename(base_name)

    doc = Document(word_file)
    image_dict = {}
    pattern = re.compile(r"(.+?)_(\d+)\.png", re.IGNORECASE)

    for file in os.listdir(output_folder):
        match = pattern.match(file)
        if match:
            placeholder, index = match.groups()
            image_dict.setdefault(placeholder, []).append((int(index), file))

    for key in image_dict:
        image_dict[key].sort()

    for para in doc.paragraphs:
        for placeholder in list(image_dict.keys()):
            if f"{{{placeholder}}}" in para.text:
                para.text = para.text.replace(f"{{{placeholder}}}", "")
                run = para.add_run(f"{{{placeholder}}}")
                run.font.color.rgb = RGBColor(255, 255, 255)
                for i, (_, image_name) in enumerate(image_dict[placeholder]):
                    image_path = os.path.join(output_folder, image_name)
                    if os.path.exists(image_path):
                        para.add_run().add_break()
                        para.add_run().add_picture(image_path, width=Cm(15), height=None)
                        if i < len(image_dict[placeholder]) - 1:
                            para.add_run().add_break(WD_BREAK.PAGE)

    doc.save(output_file)
    messagebox.showinfo("Succes", f"Dokument gemt som {output_file}")


def start_process():
    save_excel_sheets_as_png()
    root.after(100, insert_images_into_word)
    messagebox.showinfo("Succes", "Processen er fuldført!")  # Viser beskeden først
    start_button.config(text="Afslut", command=close_application)


def close_application():
    root.quit()
    root.destroy()


root = tk.Tk()
root.title("Excel til Word Automation")
root.geometry("600x800")

crop_frame = tk.LabelFrame(root, text="Beskæring (mm)", padx=10, pady=10)
crop_frame.pack(pady=10)


def create_crop_input(label_text):
    frame = tk.Frame(crop_frame)
    frame.pack(anchor="w")
    label = tk.Label(frame, text=label_text, width=12, anchor="w")
    label.pack(side="left")
    entry = tk.Entry(frame, width=10)
    entry.pack(side="left")
    return entry


top_crop_entry = create_crop_input("Top")
bottom_crop_entry = create_crop_input("Bund")
left_crop_entry = create_crop_input("Venstre")
right_crop_entry = create_crop_input("Højre")

frame = tk.Frame(root)
frame.pack(pady=10)
excel_entry = tk.Entry(frame, width=50)
excel_entry.pack(side=tk.LEFT)
tk.Button(frame, text="Vælg Excel", command=select_excel_file).pack(side=tk.LEFT)

sheets_frame = tk.Frame(root)
sheets_frame.pack(pady=10, fill="both", expand=True)

frame2 = tk.Frame(root)
frame2.pack(pady=10)
word_entry = tk.Entry(frame2, width=50)
word_entry.pack(side=tk.LEFT)
tk.Button(frame2, text="Vælg Word", command=select_word_file).pack(side=tk.LEFT)

frame3 = tk.Frame(root)
frame3.pack(pady=10)
output_entry = tk.Entry(frame3, width=50)
output_entry.pack(side=tk.LEFT)
tk.Button(frame3, text="Vælg mappe til billeder", command=select_image_output_folder).pack(side=tk.LEFT)

progress_var = tk.IntVar()
progress_bar = ttk.Progressbar(root, variable=progress_var, maximum=100)
progress_bar.pack(pady=10)
progress_label = tk.Label(root, text="")
progress_label.pack()

start_button = tk.Button(root, text="Start", command=start_process)
start_button.pack(pady=10)

root.mainloop()

