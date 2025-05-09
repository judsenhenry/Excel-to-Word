import tempfile
import streamlit as st
import xlwings as xw
from docx import Document
from docx.shared import Cm, RGBColor
from docx.enum.text import WD_BREAK
import re
import time
from PIL import Image
import os
import sys

# Find poppler-mappen, uanset om scriptet kører som .py eller .exe
if getattr(sys, 'frozen', False):
    script_dir = sys._MEIPASS  # PyInstaller bruger _MEIPASS som midlertidig mappe
else:
    script_dir = os.path.dirname(os.path.abspath(sys.argv[0])) if "__file__" in globals() else os.getcwd()

poppler_path = os.path.join(script_dir, "poppler", "Library", "bin")

# Brug Poppler-stien i pdf2image
from pdf2image import convert_from_path


def select_excel_file():
    file_path = st.file_uploader("Vælg Excel-fil", type=["xlsx", "xlsm"])
    if file_path:
        load_sheets(file_path)


def load_sheets(file_path):
    try:
        with xw.App(visible=False) as app:
            wb = app.books.open(file_path)
            sheet_names = [sheet.name for sheet in wb.sheets if sheet.visible]

        # Vis arkene som checkbokse
        selected_sheets = st.multiselect("Vælg de ark, du vil bruge", sheet_names)

        return selected_sheets
    except Exception as e:
        st.error(f"Kunne ikke indlæse ark: {e}")
        return []


def select_word_file():
    return st.file_uploader("Vælg Word-fil", type=["docx"])


def select_image_output_folder():
    return st.text_input("Vælg mappe til output")


def generate_unique_filename(base_path):
    counter = 1
    file_path = f"{base_path}_opdateret.docx"
    while os.path.exists(file_path):
        file_path = f"{base_path}_opdateret_{counter}.docx"
        counter += 1
    return file_path


def crop_image(image_path):
    try:
        img = Image.open(image_path)
        dpi = 300  # Samme dpi som ved konvertering

        # Læs værdier fra input
        top_crop = int(float(st.text_input("Top beskæring (mm)") or 0) / 25.4 * dpi)
        bottom_crop = int(float(st.text_input("Bund beskæring (mm)") or 0) / 25.4 * dpi)
        left_crop = int(float(st.text_input("Venstre beskæring (mm)") or 0) / 25.4 * dpi)
        right_crop = int(float(st.text_input("Højre beskæring (mm)") or 0) / 25.4 * dpi)

        width, height = img.size
        cropped_img = img.crop((left_crop, top_crop, width - right_crop, height - bottom_crop))
        cropped_img.save(image_path)
    except Exception as e:
        st.error(f"Beskæring fejlede for {os.path.basename(image_path)}:\n{e}")


def save_excel_sheets_as_png():
    excel_file = st.session_state.get('excel_file')
    output_folder = st.session_state.get('output_folder')
    selected_sheets = st.session_state.get('selected_sheets')

    if not excel_file or not selected_sheets or not output_folder:
        st.warning("Vælg Excel-fil, ark og billedoutput-mappe!")
        return

    try:
        with xw.App(visible=False) as app:
            wb = app.books.open(os.path.abspath(excel_file))
            total_sheets = len(selected_sheets)
            start_time = time.time()

            for index, sheet_name in enumerate(selected_sheets, start=1):
                sheet = wb.sheets[sheet_name]
                temp_pdf = os.path.join(tempfile.gettempdir(), f"{sheet_name}.pdf")

                try:
                    sheet.to_pdf(temp_pdf)
                    images = convert_from_path(temp_pdf, dpi=300, poppler_path=poppler_path)

                    for i, img in enumerate(images, start=1):
                        img_path = os.path.join(output_folder, f"{sheet_name}_{i}.png")
                        img.save(img_path, "PNG")
                        crop_image(img_path)  # Beskær billedet
                except Exception as e:
                    st.error(f"Kunne ikke eksportere {sheet_name}: {e}")

                elapsed = time.time() - start_time
                remaining = (elapsed / index) * (total_sheets - index)
                st.progress(int((index / total_sheets) * 100))
                st.text(f"Forventet tid tilbage: {int(remaining)} sekunder")

            st.success("Billeder gemt!")
    except Exception as e:
        st.error(f"Kunne ikke åbne Excel: {e}")


def insert_images_into_word():
    word_file = st.session_state.get('word_file')
    output_folder = st.session_state.get('output_folder')

    if not word_file or not output_folder:
        st.warning("Vælg Word-fil og billedoutput-mappe!")
        return

    base_name, _ = os.path.splitext(word_file)
    output_file = generate_unique_filename(base_name)

    doc = Document(word_file)
    image_dict = {}
    pattern = re.compile(r"(.+?)_(\d+)\.png", re.IGNORECASE)

    for file in os.listdir(output_folder):
        match = pattern.match(file)
        if match:
            placeholder, index = match.groups()
            image_dict.setdefault(placeholder, []).append((int(index), file))

    for key in image_dict:
        image_dict[key].sort()

    for para in doc.paragraphs:
        for placeholder in list(image_dict.keys()):
            if f"{{{placeholder}}}" in para.text:
                para.text = para.text.replace(f"{{{placeholder}}}", "")
                run = para.add_run(f"{{{placeholder}}}")
                run.font.color.rgb = RGBColor(255, 255, 255)
                for i, (_, image_name) in enumerate(image_dict[placeholder]):
                    image_path = os.path.join(output_folder, image_name)
                    if os.path.exists(image_path):
                        para.add_run().add_break()
                        para.add_run().add_picture(image_path, width=Cm(15), height=None)
                        if i < len(image_dict[placeholder]) - 1:
                            para.add_run().add_break(WD_BREAK.PAGE)

    doc.save(output_file)
    st.success(f"Dokument gemt som {output_file}")


def start_process():
    save_excel_sheets_as_png()
    insert_images_into_word()
    st.success("Processen er fuldført!")  # Viser beskeden først


st.title("Excel til Word Automation")

# Upload filer og vælg mapper
select_excel_file()
selected_sheets = load_sheets(st.session_state.get('excel_file'))  # Load selected sheets
select_word_file()
select_image_output_folder()

if st.button("Start"):
    start_process()
